from langchain.schema import Document
from PyPDF2 import PdfReader
import docx
import textract
from bs4 import BeautifulSoup
from io import BytesIO

# -----------------------
# File converters (in-memory)
# -----------------------
def pdf_to_text(file_content):
    reader = PdfReader(file_content if isinstance(file_content, BytesIO) else BytesIO(file_content))
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def docx_to_text(file_content):
    doc = docx.Document(file_content if isinstance(file_content, BytesIO) else BytesIO(file_content))
    return "\n".join([para.text for para in doc.paragraphs])

def doc_to_text(file_content):
    return textract.process(BytesIO(file_content) if not isinstance(file_content, (bytes, BytesIO)) else file_content).decode("utf-8", errors="ignore")

def txt_to_text(file_content):
    if isinstance(file_content, bytes):
        return file_content.decode("utf-8", errors="ignore")
    elif isinstance(file_content, BytesIO):
        return file_content.read().decode("utf-8", errors="ignore")
    else:
        return str(file_content)

def html_to_text(file_content):
    if isinstance(file_content, bytes):
        html = file_content.decode("utf-8", errors="ignore")
    elif isinstance(file_content, BytesIO):
        html = file_content.read().decode("utf-8", errors="ignore")
    else:
        html = str(file_content)

    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    return "\n".join([line.strip() for line in text.splitlines() if line.strip()])


def convert_documents_to_langchain_docs(document_objs: list) -> list[Document]:
    lc_docs = []

    for doc_obj in document_objs:
        ext = doc_obj.filename.split('.')[-1].lower()
        file_content = getattr(doc_obj, "file_content", None)
        if not file_content:
            raise ValueError(f"No file content provided for {doc_obj.filename}")

        try:
            if ext == "pdf":
                text = pdf_to_text(file_content)
            elif ext == "docx":
                text = docx_to_text(file_content)
            elif ext == "doc":
                text = doc_to_text(file_content)
            elif ext == "txt":
                text = txt_to_text(file_content)
            elif ext in ["html", "htm"]:
                text = html_to_text(file_content)
            else:
                raise ValueError(f"Unsupported file type: {ext}")

            lc_docs.append(Document(page_content=text, metadata={"source": doc_obj.filename}))
        except Exception as e:
            print(f"Skipping {doc_obj.filename} due to error: {e}")

    return lc_docs
